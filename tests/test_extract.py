from io import BytesIO

from src.documents.extract import extract_text_preview


def tiny_pdf(text: str | list[str]) -> bytes:
    texts = [text] if isinstance(text, str) else text
    n_pages = len(texts)
    font_obj_num = 3 + 2 * n_pages
    page_obj_nums = [3 + 2 * i for i in range(n_pages)]
    objects = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        (
            f"<</Type/Pages/Kids[{' '.join(f'{n} 0 R' for n in page_obj_nums)}]"
            f"/Count {n_pages}>>"
        ).encode(),
    ]
    for page_text in texts:
        stream = f"BT /F1 12 Tf 72 720 Td ({page_text}) Tj ET".encode()
        objects.append(
            b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents "
            + str(len(objects) + 1 + 1).encode()
            + b" 0 R/Resources<</Font<</F1 "
            + str(font_obj_num).encode()
            + b" 0 R>>>>>>"
        )
        objects.append(
            b"<</Length " + str(len(stream)).encode() + b">>stream\n" + stream + b"\nendstream"
        )
    objects.append(b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>")

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj".encode() + obj + b"endobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer<</Size {len(objects) + 1}/Root 1 0 R>>\nstartxref\n{xref_pos}\n%%EOF".encode()
    return bytes(out)


def tiny_docx(text: str) -> bytes:
    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(buf)
    return buf.getvalue()


def test_pdf_preview_extracts_text():
    preview = extract_text_preview(tiny_pdf("Master Services Agreement v2"), "msa.pdf")
    assert "Master Services Agreement v2" in preview


def test_docx_preview_extracts_text():
    preview = extract_text_preview(tiny_docx("Section 4.2 Limitation of Liability"), "msa.docx")
    assert "Section 4.2 Limitation of Liability" in preview


def test_preview_respects_max_chars():
    preview = extract_text_preview(tiny_docx("word " * 2000), "long.docx", max_chars=100)
    assert len(preview) == 100


def test_unsupported_extension_returns_empty():
    assert extract_text_preview(b"anything", "photo.png") == ""


def test_corrupt_pdf_returns_empty_not_error():
    assert extract_text_preview(b"%PDF-not really", "broken.pdf") == ""


def test_pdf_preview_reads_beyond_three_pages_for_large_max_chars():
    pages = [f"Page {i} marker text" for i in range(1, 6)]
    preview = extract_text_preview(tiny_pdf(pages), "msa.pdf", max_chars=10_000)
    assert "Page 4 marker text" in preview


def test_docx_preserves_paragraph_breaks():
    from tests.test_extract import tiny_docx  # noqa: F401  (module self-import guard)
    from docx import Document as DocxDocument
    buf = BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("1. TERM. The   term is twelve months.")
    doc.add_paragraph("2. LIABILITY. Liability is unlimited.")
    doc.save(buf)
    preview = extract_text_preview(buf.getvalue(), "msa.docx")
    assert "1. TERM. The term is twelve months.\n\n2. LIABILITY. Liability is unlimited." == preview


def test_pdf_splits_numbered_clauses_into_paragraphs():
    pages = [["INTRO AGREEMENT TITLE", "1. TERM. Twelve months", "continued on same clause.", "2. FEES. Ninety days."]]
    # one page, four layout lines
    content = tiny_pdf(["INTRO AGREEMENT TITLE\n1. TERM. Twelve months\ncontinued on same clause.\n2. FEES. Ninety days."])
    preview = extract_text_preview(content, "msa.pdf")
    assert "INTRO AGREEMENT TITLE" in preview
    assert "\n\n1. TERM. Twelve months continued on same clause." in preview
    assert "\n\n2. FEES. Ninety days." in preview
