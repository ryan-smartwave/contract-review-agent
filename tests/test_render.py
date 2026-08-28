from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.documents.render import render_docx


def test_render_docx_returns_openable_bytes():
    data = render_docx("Some text.")
    doc = Document(BytesIO(data))
    assert len(doc.paragraphs) == 1


def test_render_docx_centers_and_bolds_title_block():
    data = render_docx("MASTER SERVICES AGREEMENT\n\n1. TERM. Twelve months.")
    doc = Document(BytesIO(data))
    assert len(doc.paragraphs) == 2

    title, body = doc.paragraphs
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title.runs[0].bold is True

    assert "Twelve months" in body.text
    assert body.alignment != WD_ALIGN_PARAGRAPH.CENTER


def test_render_docx_skips_blank_blocks():
    data = render_docx("A\n\n\n\nB")
    doc = Document(BytesIO(data))
    assert [p.text for p in doc.paragraphs] == ["A", "B"]
