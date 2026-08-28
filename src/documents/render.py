from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

TITLE_MAX_CHARS = 80


def _is_title(block: str) -> bool:
    return (
        len(block) < TITLE_MAX_CHARS
        and any(c.isalpha() for c in block)
        and block == block.upper()
    )


def render_docx(text: str) -> bytes:
    doc = Document()
    for block in (b.strip() for b in text.split("\n\n")):
        if not block:
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(block)
        if _is_title(block):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
